from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\pactr\Downloads\VATM-AeroSync-Implemented-Format-Report.docx")
OUTPUT = Path(
    r"C:\Users\pactr\Downloads\VATM-AeroSync-Implemented-Format-Report-Updated-2026-07-28.docx"
)

OPERATORS = {
    "OF-5199 (1).docx": "RAYA AIRWAYS (ICAO RMY; IATA TH)",
    "LD-06.A.S.2026VN.REV8 (1).doc": "FEDERAL EXPRESS (ICAO FDX; IATA FX)",
    "(SPA066) REV1 LD-2631 C26_18Jul-18Jul_ND.docx": (
        "SUN PHUQUOC (ICAO SPQ; IATA 9G)"
    ),
    "cor(07feb)_LD-545.02.2025VN-rvs-ld38a-qh1123_14feb.docx": (
        "Bamboo Airways (ICAO BAV; IATA QH)"
    ),
    "OF-5277 (REV1).docx": "VISTAJET LIMITED (ICAO VJT; IATA not listed)",
    "(SPA017) REV9 LD-32B S26_23JULY-31JUL_ND.docx": (
        "SUN PHUQUOC (ICAO SPQ; IATA 9G)"
    ),
    "LD-83_A_S_2026VN.RVS.22JUL26.docx": (
        "JEJUAIR (REPUBLIC OF KOREA) (ICAO JJA; IATA 7C)"
    ),
}

FORMAT_ROWS = [
    (
        "1",
        "CAAV English scheduled overflight",
        OPERATORS["OF-5199 (1).docx"],
        ".docx",
        "Scheduled flights",
        "Implemented",
    ),
    (
        "2",
        "CAAV English landing revision",
        OPERATORS["LD-06.A.S.2026VN.REV8 (1).doc"],
        ".doc",
        "2.2 New Schedule",
        "Implemented",
    ),
    (
        "3",
        "SPA066 Vietnamese landing revision",
        OPERATORS["(SPA066) REV1 LD-2631 C26_18Jul-18Jul_ND.docx"],
        ".docx",
        "2.2 New Schedule",
        "Implemented",
    ),
    (
        "4",
        "CAAV Vietnamese landing correction",
        OPERATORS["cor(07feb)_LD-545.02.2025VN-rvs-ld38a-qh1123_14feb.docx"],
        ".docx",
        "2.2 New Schedule",
        "Implemented",
    ),
    (
        "5",
        "CAAV English overflight revision",
        OPERATORS["OF-5277 (REV1).docx"],
        ".docx",
        "3.2 New",
        "Implemented",
    ),
    (
        "6",
        "SPA017 Vietnamese seasonal revision",
        OPERATORS["(SPA017) REV9 LD-32B S26_23JULY-31JUL_ND.docx"],
        ".docx",
        "2.2 New Schedule",
        "Implemented",
    ),
    (
        "7",
        "CAAV English issued landing permit revision",
        OPERATORS["LD-83_A_S_2026VN.RVS.22JUL26.docx"],
        ".docx",
        "2.2 New Schedule",
        "Implemented",
    ),
]


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table_text(table):
    centered_columns = {0, 3, 5}
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if col_index in centered_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = None
                    run.font.color.theme_color = None
                    run.font.color.rgb = __import__(
                        "docx"
                    ).shared.RGBColor(0x0B, 0x25, 0x45)
                elif col_index == 5:
                    run.bold = True
                    run.font.color.rgb = __import__(
                        "docx"
                    ).shared.RGBColor(0x1F, 0x6D, 0x3A)


def replace_run_text(document, replacements):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for old, new in replacements:
                if old in run.text:
                    run.text = run.text.replace(old, new)


def insert_paragraph_after(paragraph, label, value, space_after=3):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p.getparent().remove(inserted._p)
    new_p.getparent().replace(new_p, inserted._p)
    inserted.style = "Normal"
    inserted.paragraph_format.space_after = Pt(space_after)
    label_run = inserted.add_run(label)
    label_run.bold = True
    inserted.add_run(value)
    return inserted


def insert_before(target, text, style="Normal", label=None, space_after=3):
    paragraph = target._parent.add_paragraph(style=style)
    paragraph._p.getparent().remove(paragraph._p)
    target._p.addprevious(paragraph._p)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if label is None:
        paragraph.add_run(text)
    else:
        label_run = paragraph.add_run(label)
        label_run.bold = True
        paragraph.add_run(text)
    return paragraph


def copy_row_format(source_row, target_row):
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        source_tc_pr = source_cell._tc.tcPr
        if source_tc_pr is not None:
            target_cell._tc.remove(target_cell._tc.tcPr)
            target_cell._tc.insert(0, deepcopy(source_tc_pr))


document = Document(SOURCE)

replace_run_text(
    document,
    [
        ("27 July 2026", "28 July 2026"),
        ("Six Word permit profiles", "Seven Word permit profiles"),
        ("six configured Word permit layouts", "seven configured Word permit layouts"),
        ("Six profile-driven Word permit formats", "Seven profile-driven Word permit formats"),
        ("six configured profiles", "seven configured profiles"),
        ("six documented Word permit families", "seven documented Word permit families"),
        ("Six Word permit profiles plus", "Seven Word permit profiles plus"),
        ("two reported mismatches", "three reported mismatches"),
        ("The two reported mismatches", "The three reported mismatches"),
    ],
)

# Add a precise reference-data note near the report metadata.
database_target = next(
    p for p in document.paragraphs if p.text.startswith("Database target:")
)
insert_paragraph_after(
    database_target,
    "Operator reference: ",
    "Oracle M_OPER (M_AERO is the aerodrome reference table).",
)

# Add operator identity to each existing profile section.
for representative, operator in OPERATORS.items():
    if representative == "LD-83_A_S_2026VN.RVS.22JUL26.docx":
        continue
    representative_paragraph = next(
        p
        for p in document.paragraphs
        if p.text == f"Representative document: {representative}"
    )
    insert_paragraph_after(
        representative_paragraph,
        "Airline/operator: ",
        operator,
    )

# Replace the format matrix with a six-column version that includes operators.
old_matrix = document.tables[0]
matrix = document.add_table(rows=1, cols=6)
matrix.style = old_matrix.style
old_matrix._tbl.addprevious(matrix._tbl)
headers = ["No.", "Format", "Airline / operator", "File", "Imported schedule", "Status"]
for cell, text in zip(matrix.rows[0].cells, headers):
    cell.text = text
    shade_cell(cell, "F2F4F7")
set_repeat_table_header(matrix.rows[0])

for row_values in FORMAT_ROWS:
    row = matrix.add_row()
    for cell, text in zip(row.cells, row_values):
        cell.text = text

set_table_geometry(matrix, [450, 2200, 2200, 650, 2450, 1410])
format_table_text(matrix)
old_matrix._element.getparent().remove(old_matrix._element)

# Insert the newest Word profile immediately before the data-extraction section.
data_extracted = next(
    p for p in document.paragraphs if p.text == "Data extracted from Word permits"
)
insert_before(
    data_extracted,
    "7. CAAV English issued landing permit revision",
    style="Heading 2",
    space_after=0,
)
insert_before(
    data_extracted,
    "caav-english-issued-permit-revision.yaml",
    label="Profile: ",
)
insert_before(
    data_extracted,
    "LD-83_A_S_2026VN.RVS.22JUL26.docx",
    label="Representative document: ",
)
insert_before(
    data_extracted,
    OPERATORS["LD-83_A_S_2026VN.RVS.22JUL26.docx"],
    label="Airline/operator: ",
)
insert_before(
    data_extracted,
    "Imports only section 2.2 New Schedule, excludes the original-permit column, "
    "allows IATA airport codes, and permits an empty airways table.",
    label="Import behavior: ",
)
insert_before(
    data_extracted,
    "LD 0083A/S/CHK/2026",
    label="Normalized permit example: ",
)
insert_before(
    data_extracted,
    "Aircraft aliases 738/7M8, 7M8/738, 738, and B738 map to CRAFT_ID 249, "
    "MTOW 0, purpose PAX.",
    label="Configured mapping: ",
)
insert_before(
    data_extracted,
    "The synthetic profile fixture passes. The supplied real document is recognized, "
    "but the current parser selects related permit LD-230 instead of primary permit "
    "LD-83; correction is required before acceptance.",
    label="Verification: ",
    space_after=9,
)

# Add an explicit verification row for the newly implemented profile.
verification_table = document.tables[2]
new_row = verification_table.add_row()
copy_row_format(verification_table.rows[1], new_row)
verification_values = [
    "LD-83 / Jeju Air profile",
    "2",
    "1",
    "Synthetic fixture passes; the real file is recognized but selects related permit "
    "LD-230 instead of primary permit LD-83.",
    "Review required",
]
for index, (cell, text) in enumerate(zip(new_row.cells, verification_values)):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if index in {1, 2, 4} else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.size = Pt(8.5)
        if index == 4:
            run.bold = True

document.save(OUTPUT)
print(OUTPUT)
